import os
import re
from typing import Tuple, List
from app.core.logging import logger


class CVParserService:
    """Servicio para lectura y extracción completa de texto e información de archivos CV (PDF, DOCX, TXT)."""

    COMMON_SKILLS = [
        # Backend
        "python", "javascript", "typescript", "java", "c#", "c++", "php", "ruby", "go", "rust",
        # Frameworks
        "fastapi", "django", "flask", "spring", "laravel", "express", "node.js", "react",
        "vue", "angular", "next.js", "nest.js", ".net",
        # DB
        "sql", "postgresql", "mysql", "mongodb", "redis", "sqlite", "oracle", "elasticsearch",
        # Cloud / DevOps
        "docker", "kubernetes", "aws", "gcp", "azure", "ci/cd", "jenkins", "github actions",
        "terraform", "linux", "bash", "nginx",
        # Data / IA
        "pandas", "numpy", "scikit-learn", "tensorflow", "pytorch", "machine learning",
        "deep learning", "power bi", "tableau", "seaborn", "matplotlib",
        # Herramientas
        "git", "jira", "confluence", "postman", "figma", "excel", "html", "css", "tailwind",
        # Soft skills / idiomas
        "english", "ingles", "español", "comunicación", "liderazgo", "trabajo en equipo",
        "scrum", "agile",
        # Sysadmin / Redes
        "cisco", "ccna", "active directory", "sysadmin", "vmware", "redes", "networking",
        "ciberseguridad", "cybersecurity", "seiq", "splunk",
    ]

    @classmethod
    def parse_cv_file(cls, file_path: str) -> Tuple[str, List[str]]:
        """Extrae el contenido completo en texto plano y habilidades de un archivo CV."""
        ext = os.path.splitext(file_path)[1].lower()
        extracted_text = ""

        try:
            if ext == ".pdf":
                extracted_text = cls._parse_pdf(file_path)
            elif ext in [".docx", ".doc"]:
                extracted_text = cls._parse_docx(file_path)
            elif ext in [".txt", ".md"]:
                extracted_text = cls._parse_txt(file_path)
            else:
                logger.warning(f"Formato no soportado para lectura directa: {ext}")
                extracted_text = cls._parse_txt(file_path)
        except Exception as e:
            logger.error(f"Error parseando el archivo CV {file_path}: {e}")
            extracted_text = f"Error al procesar el archivo CV: {str(e)}"

        skills = cls.extract_skills(extracted_text)
        return extracted_text.strip(), skills

    @staticmethod
    def _parse_pdf(file_path: str) -> str:
        """
        Extrae texto de un PDF priorizando PyMuPDF (fitz), luego pdfplumber, luego pypdf.
        El fallback raw se eliminó para evitar extraer binarios del PDF.
        """
        text_pages = []

        # ── 1. PyMuPDF (fitz) — el más preciso para PDFs complejos (Canva, LaTeX, etc.) ──
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(file_path)
            for page_num, page in enumerate(doc):
                # "text" mode: texto en orden de lectura natural
                t = page.get_text("text")
                if t and t.strip():
                    text_pages.append(t.strip())

            doc.close()

            if text_pages:
                full_text = "\n\n".join(text_pages)
                # Verificar que realmente tiene texto legible (no solo símbolos o basura)
                legible_words = re.findall(r'[a-zA-ZáéíóúÁÉÍÓÚñÑüÜ]{3,}', full_text)
                if len(legible_words) >= 10:
                    logger.info(f"PyMuPDF extrajo {len(legible_words)} palabras del PDF")
                    return full_text

        except ImportError:
            logger.warning("PyMuPDF no está instalado. Instalá: pip install pymupdf")
        except Exception as e:
            logger.warning(f"PyMuPDF error: {e}")

        # ── 2. pdfplumber — excelente para tablas y columnas ──
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text_pages.append(page_text.strip())

            if text_pages:
                full_text = "\n\n".join(text_pages)
                legible_words = re.findall(r'[a-zA-ZáéíóúÁÉÍÓÚñÑ]{3,}', full_text)
                if len(legible_words) >= 10:
                    logger.info(f"pdfplumber extrajo {len(legible_words)} palabras del PDF")
                    return full_text

        except ImportError:
            logger.warning("pdfplumber no está instalado. Instalá: pip install pdfplumber")
        except Exception as e:
            logger.warning(f"pdfplumber error: {e}")

        # ── 3. pypdf — fallback estándar ──
        try:
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            for page in reader.pages:
                try:
                    page_text = page.extract_text(extraction_mode="layout")
                except Exception:
                    page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_pages.append(page_text.strip())

            if text_pages:
                full_text = "\n\n".join(text_pages)
                legible_words = re.findall(r'[a-zA-ZáéíóúÁÉÍÓÚñÑ]{3,}', full_text)
                if len(legible_words) >= 10:
                    logger.info(f"pypdf extrajo {len(legible_words)} palabras del PDF")
                    return full_text

        except ImportError:
            logger.warning("pypdf no está instalado. Instalá: pip install pypdf")
        except Exception as e:
            logger.warning(f"pypdf error: {e}")

        # ── 4. Último recurso: extraer texto desde struct tree del PDF (para PDFs tagged como Canva) ──
        try:
            # Los PDFs de Canva tienen StructTree con /E (alternate text) que contiene texto real
            with open(file_path, "rb") as f:
                raw = f.read().decode("latin-1", errors="ignore")

            # Extraer contenido de los campos /E (alt text en StructTree) que son texto real
            e_matches = re.findall(r'/E\s+([^\n/]+?)(?:\s*/)', raw)
            t_matches = re.findall(r'/T\s+([^\n/]+?)(?:\s*/)', raw)

            all_text_fragments = []
            for m in (e_matches + t_matches):
                m = m.strip()
                # Solo incluir si tiene caracteres legibles (al menos 4 letras consecutivas)
                if re.search(r'[a-zA-ZáéíóúñÑ]{4,}', m):
                    all_text_fragments.append(m)

            if all_text_fragments:
                result = "\n".join(dict.fromkeys(all_text_fragments))  # deduplicar
                legible_words = re.findall(r'[a-zA-ZáéíóúÁÉÍÓÚñÑ]{3,}', result)
                if len(legible_words) >= 5:
                    logger.info(f"StructTree extrajo {len(legible_words)} palabras del PDF Canva")
                    return result

        except Exception as e:
            logger.error(f"Error en extracción StructTree: {e}")

        logger.error(f"No se pudo extraer texto legible del PDF: {file_path}. "
                     "Instalá pymupdf: pip install pymupdf")
        return (
            "⚠️ No se pudo extraer texto del PDF.\n"
            "El sistema probó PyMuPDF, pdfplumber y pypdf sin éxito.\n"
            "Por favor, pegá el texto de tu CV manualmente en este campo."
        )

    @staticmethod
    def _parse_docx(file_path: str) -> str:
        try:
            import docx
            doc = docx.Document(file_path)
            full_text = []

            for para in doc.paragraphs:
                if para.text and para.text.strip():
                    full_text.append(para.text.strip())

            for table in doc.tables:
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                    if row_cells:
                        unique_cells = []
                        for cell in row_cells:
                            if not unique_cells or unique_cells[-1] != cell:
                                unique_cells.append(cell)
                        full_text.append(" | ".join(unique_cells))

            return "\n".join(full_text)
        except ImportError:
            logger.warning("python-docx no está instalado.")
            return "Contenido de CV DOCX (instalar python-docx para lectura completa)."
        except Exception as e:
            logger.error(f"Error procesando DOCX: {e}")
            return ""

    @staticmethod
    def _parse_txt(file_path: str) -> str:
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception as e:
            logger.error(f"Error procesando TXT: {e}")
            return ""

    @classmethod
    def extract_skills(cls, text: str) -> List[str]:
        if not text:
            return []
        lower_text = text.lower()
        found_skills = set()
        for skill in cls.COMMON_SKILLS:
            # Búsqueda flexible: con o sin límite de palabra para skills multi-palabra
            if " " in skill:
                if skill in lower_text:
                    found_skills.add(skill.title())
            else:
                pattern = r'\b' + re.escape(skill) + r'\b'
                if re.search(pattern, lower_text):
                    found_skills.add(skill.capitalize() if len(skill) > 3 else skill.upper())
        return sorted(list(found_skills))
